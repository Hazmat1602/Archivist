from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docxtpl import DocxTemplate
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.auth import get_current_user
from app.database import get_db
from app.models.box import Box
from app.models.folder import Folder
from app.models.user import User

router = APIRouter(prefix="/labels", tags=["labels"])

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _format_date(value: date | datetime | str | None) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, (date, datetime)):
        return value.strftime("%d-%m-%Y")
    return ""


def _render_label_pages(
    template_path: Path,
    contexts: list[dict[str, str]],
) -> BytesIO:
    """Render one or more pages of labels using a docx template."""
    rendered_document: Document | None = None

    for context in contexts:
        template = DocxTemplate(template_path)
        template.render(context)
        buffer = BytesIO()
        template.save(buffer)
        buffer.seek(0)

        if rendered_document is None:
            rendered_document = Document(buffer)
            continue

        next_page = Document(buffer)
        for element in next_page.element.body:
            rendered_document.element.body.append(element)

    output = BytesIO()
    if rendered_document is None:
        rendered_document = Document()
    rendered_document.save(output)
    output.seek(0)
    return output


# ---------------------------------------------------------------------------
# Folder labels
# ---------------------------------------------------------------------------

FOLDER_LABELS_PER_SHEET = 24


class FolderLabelRequest(BaseModel):
    folder_ids: Optional[list[int]] = None


def _folder_label_contexts(
    folders: list[Folder],
) -> list[dict[str, str]]:
    contexts: list[dict[str, str]] = []

    for offset in range(0, len(folders), FOLDER_LABELS_PER_SHEET):
        chunk = folders[offset : offset + FOLDER_LABELS_PER_SHEET]
        context: dict[str, str] = {}

        for index in range(FOLDER_LABELS_PER_SHEET):
            number = index + 1
            folder = chunk[index] if index < len(chunk) else None

            context[f"code_{number}"] = (
                folder.retention_id if folder else ""
            )
            context[f"name_{number}"] = folder.name if folder else ""
            context[f"create_{number}"] = (
                _format_date(folder.created_date) if folder else ""
            )
            context[f"start_{number}"] = (
                _format_date(folder.start_date) if folder else ""
            )
            context[f"end_{number}"] = (
                _format_date(folder.expiry_date) if folder else ""
            )

        contexts.append(context)

    if not contexts:
        empty: dict[str, str] = {}
        for number in range(1, FOLDER_LABELS_PER_SHEET + 1):
            empty[f"code_{number}"] = ""
            empty[f"name_{number}"] = ""
            empty[f"create_{number}"] = ""
            empty[f"start_{number}"] = ""
            empty[f"end_{number}"] = ""
        contexts.append(empty)

    return contexts


@router.post("/folders")
def create_folder_labels(
    body: FolderLabelRequest,
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    template_path = TEMPLATES_DIR / "folder_labels_template.docx"
    if not template_path.exists():
        raise HTTPException(500, "Folder label template not found")

    if body.folder_ids:
        folders = (
            db.query(Folder)
            .filter(Folder.id.in_(body.folder_ids))
            .all()
        )
        # Preserve the order the user selected
        order = {fid: idx for idx, fid in enumerate(body.folder_ids)}
        folders.sort(key=lambda f: order.get(f.id, 0))
    else:
        folders = db.query(Folder).order_by(Folder.id).all()

    contexts = _folder_label_contexts(folders)
    document_bytes = _render_label_pages(template_path, contexts)

    return StreamingResponse(
        document_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="folder_labels.docx"'
        },
    )


# ---------------------------------------------------------------------------
# Box labels
# ---------------------------------------------------------------------------

BOX_LABELS_PER_SHEET = 4
MAX_FOLDER_LINES = 20


class BoxLabelRequest(BaseModel):
    box_ids: Optional[list[int]] = None


def _get_assigned_folders(
    db: Session,
    box_id: int,
) -> tuple[str, str]:
    """Return two formatted columns of folder retention IDs for a box."""
    folder_ids = [
        rid
        for (rid,) in (
            db.query(Folder.retention_id)
            .filter(Folder.box_id == box_id)
            .order_by(Folder.id)
        )
    ]

    def format_lines(ids: list[str]) -> str:
        padded = ids + ["" for _ in range(max(0, MAX_FOLDER_LINES - len(ids)))]
        return "\n".join(f"\u2022    {fid}" for fid in padded)

    column_a = format_lines(folder_ids[:MAX_FOLDER_LINES])
    column_b = format_lines(folder_ids[MAX_FOLDER_LINES : MAX_FOLDER_LINES * 2])

    return column_a, column_b


def _box_label_contexts(
    boxes: list[Box],
    db: Session,
) -> list[dict[str, str]]:
    contexts: list[dict[str, str]] = []

    for offset in range(0, len(boxes), BOX_LABELS_PER_SHEET):
        chunk = boxes[offset : offset + BOX_LABELS_PER_SHEET]
        context: dict[str, str] = {}

        for index in range(BOX_LABELS_PER_SHEET):
            number = index + 1
            box = chunk[index] if index < len(chunk) else None

            col_a, col_b = (
                _get_assigned_folders(db, box.id) if box else ("", "")
            )

            context[f"code_{number}"] = box.code if box else ""
            context[f"name_{number}"] = (box.name or "") if box else ""
            context[f"create_{number}"] = (
                _format_date(box.created_date) if box else ""
            )
            context[f"end_{number}"] = (
                _format_date(box.expiry_date) if box else ""
            )
            context[f"column_a_{number}"] = col_a
            context[f"column_b_{number}"] = col_b

        contexts.append(context)

    if not contexts:
        empty: dict[str, str] = {}
        for number in range(1, BOX_LABELS_PER_SHEET + 1):
            empty[f"code_{number}"] = ""
            empty[f"name_{number}"] = ""
            empty[f"create_{number}"] = ""
            empty[f"end_{number}"] = ""
            empty[f"column_a_{number}"] = ""
            empty[f"column_b_{number}"] = ""
        contexts.append(empty)

    return contexts


@router.post("/boxes")
def create_box_labels(
    body: BoxLabelRequest,
    db: Session = Depends(get_db),
    _user: User = Depends(get_current_user),
):
    template_path = TEMPLATES_DIR / "box_label_template.docx"
    if not template_path.exists():
        raise HTTPException(500, "Box label template not found")

    if body.box_ids:
        boxes = (
            db.query(Box).filter(Box.id.in_(body.box_ids)).all()
        )
        order = {bid: idx for idx, bid in enumerate(body.box_ids)}
        boxes.sort(key=lambda b: order.get(b.id, 0))
    else:
        boxes = db.query(Box).order_by(Box.id).all()

    contexts = _box_label_contexts(boxes, db)
    document_bytes = _render_label_pages(template_path, contexts)

    return StreamingResponse(
        document_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="box_labels.docx"'
        },
    )
